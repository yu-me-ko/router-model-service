import csv
import re
import sys
from pathlib import Path

INPUT_PATH = Path("data/raw/华工学生常见问题收集表_2000条.docx")
OUTPUT_PATH = Path("data/extracted_public_qa.csv")
FIELDNAMES = ["question", "workflow", "tools", "sourceCategory", "keywords"]

TIME_KEYWORDS = [
    "今天",
    "明天",
    "后天",
    "周末",
    "下周",
    "几点",
    "什么时候",
    "时间",
    "开放",
    "开门",
    "关门",
    "上班",
]

CATEGORY_RE = re.compile(r"^【(.+?)】$")
QUESTION_RE = re.compile(r"^[★☆\s]*\d{1,5}\s*[.．、)]\s*(.+)$")
KEYWORDS_RE = re.compile(r"^关键词[:：]\s*(.+)$")


def iter_docx_lines(path: Path):
    from docx import Document

    document = Document(path)

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            yield text

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        yield text


def clean_question(text: str) -> str:
    text = text.strip()
    text = re.sub(r"^[★☆\s]+", "", text)
    text = re.sub(r"^\d{1,5}\s*[.．、)]\s*", "", text)
    return re.sub(r"\s+", " ", text).strip()


def infer_tools(question: str) -> str:
    tools = ["SEARCH_PUBLIC_KNOWLEDGE"]
    if any(keyword in question for keyword in TIME_KEYWORDS):
        tools.insert(0, "GET_CURRENT_TIME")
    return "|".join(tools)


def extract_rows(path: Path) -> list[dict[str, str]]:
    rows = []
    seen_questions = set()
    current_category = ""
    pending_row = None

    def flush_pending():
        nonlocal pending_row
        if pending_row is not None:
            rows.append(pending_row)
            pending_row = None

    for line in iter_docx_lines(path):
        category_match = CATEGORY_RE.match(line)
        if category_match:
            flush_pending()
            current_category = category_match.group(1).strip()
            continue

        keywords_match = KEYWORDS_RE.match(line)
        if keywords_match:
            if pending_row is not None:
                pending_row["keywords"] = keywords_match.group(1).strip()
                flush_pending()
            continue

        question_match = QUESTION_RE.match(line)
        if not question_match:
            continue

        flush_pending()
        question = clean_question(question_match.group(1))
        if not question or question in seen_questions:
            continue

        seen_questions.add(question)
        pending_row = {
            "question": question,
            "workflow": "KNOWLEDGE_QA",
            "tools": infer_tools(question),
            "sourceCategory": current_category,
            "keywords": "",
        }

    flush_pending()
    return rows


def write_rows(rows: list[dict[str, str]]) -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    temp_output = OUTPUT_PATH.with_suffix(".tmp")

    with open(temp_output, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=FIELDNAMES)
        writer.writeheader()
        writer.writerows(rows)

    temp_output.replace(OUTPUT_PATH)


def main() -> None:
    if not INPUT_PATH.exists():
        print("请将 docx 文件放入 data/raw/ 目录。")
        print(f"期望文件路径: {INPUT_PATH}")
        sys.exit(1)

    try:
        rows = extract_rows(INPUT_PATH)
    except ModuleNotFoundError as exc:
        if exc.name == "docx":
            print("缺少依赖 python-docx，请先安装：pip install -r requirements.txt")
            sys.exit(1)
        raise

    write_rows(rows)

    print("提取完成")
    print("总条数:", len(rows))
    print("输出文件:", OUTPUT_PATH)


if __name__ == "__main__":
    main()
